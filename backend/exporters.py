"""
Export Utilities — AI Tutor Golden Dataset Exporter
Generates PDFs and DOCX files matching the exact reference Golden Dataset format.
"""
import io
import os
import json
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Register Segoe UI Emoji or fallback font if available
EMOJI_FONT = None
try:
    emoji_path = r'C:\Windows\Fonts\seguiemj.ttf'
    if os.path.exists(emoji_path):
        pdfmetrics.registerFont(TTFont('SegoeUIEmoji', emoji_path))
        EMOJI_FONT = 'SegoeUIEmoji'
except Exception as e:
    print(f"[Exporters] Emoji font registration notice: {e}")


def dataset_to_pdf(dataset: dict) -> bytes:
    """Convert golden dataset dict to PDF bytes matching exact Golden Dataset reference format."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=14.4,
        rightMargin=14.4,
        topMargin=18,
        bottomMargin=18
    )
    
    styles = getSampleStyleSheet()

    # Typography & Styles
    s_top_title = ParagraphStyle(
        'TopTitle', parent=styles['Normal'],
        fontName='Helvetica-Bold', fontSize=16.5, leading=20,
        textColor=colors.HexColor('#1a2b4c'), alignment=1, spaceAfter=8
    )

    s_meta_text = ParagraphStyle(
        'MetaText', parent=styles['Normal'],
        fontName='Times-Roman', fontSize=8.1, leading=12,
        textColor=colors.HexColor('#1a1a1a')
    )

    s_subtopic_title = ParagraphStyle(
        'SubtopicTitle', parent=styles['Normal'],
        fontName='Helvetica-Bold', fontSize=13.3, leading=16,
        textColor=colors.HexColor('#1a2b4c'), spaceBefore=10, spaceAfter=4,
        keepWithNext=True
    )

    s_banner = ParagraphStyle(
        'BannerText', parent=styles['Normal'],
        fontName='Helvetica-Bold', fontSize=8.6, leading=11,
        textColor=colors.white
    )

    s_body = ParagraphStyle(
        'BodyText', parent=styles['Normal'],
        fontName='Times-Roman', fontSize=8.1, leading=11.5,
        textColor=colors.HexColor('#1a1a1a'), spaceBefore=2, spaceAfter=4
    )

    s_bullet = ParagraphStyle(
        'BulletText', parent=styles['Normal'],
        fontName='Times-Roman', fontSize=8.1, leading=11.5,
        textColor=colors.HexColor('#1a1a1a'), leftIndent=12, spaceBefore=1, spaceAfter=2
    )

    story = []

    # 1. Top Title
    story.append(Paragraph('AI Tutor Study Module', s_top_title))

    # 2. Metadata Box
    meta = dataset.get('metadata', {})
    book_name = meta.get('book_name', 'BASIC CONCEPTS OF COMPANY & ITS STRUCTURE')
    chapter_name = meta.get('chapter_name', 'Chapter 1 Concepts of "company"')
    chapter_topic = meta.get('chapter_topic', '')
    sub_topics_covered = meta.get('sub_topics_covered', [])
    if isinstance(sub_topics_covered, list):
        sub_topics_str = " | ".join([str(st) for st in sub_topics_covered])
    else:
        sub_topics_str = str(sub_topics_covered)

    meta_html = (
        f"<b>Book Name:</b> {book_name}<br/>"
        f"<b>Chapter Name:</b> {chapter_name}<br/>"
        f"<b>Chapter Topic:</b> {chapter_topic}<br/>"
        f"<b>Sub Topics Covered:</b> {sub_topics_str}"
    )
    meta_p = Paragraph(meta_html, s_meta_text)

    t_meta = Table([[ '', meta_p ]], colWidths=[3, 563.47])
    t_meta.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (0,0), colors.HexColor('#1a2b4c')),
        ('BACKGROUND', (1,0), (1,0), colors.HexColor('#f2f5fa')),
        ('LEFTPADDING', (1,0), (1,0), 8),
        ('RIGHTPADDING', (1,0), (1,0), 8),
        ('TOPPADDING', (1,0), (1,0), 6),
        ('BOTTOMPADDING', (1,0), (1,0), 6),
        ('LEFTPADDING', (0,0), (0,0), 0),
        ('RIGHTPADDING', (0,0), (0,0), 0),
    ]))
    story.append(t_meta)
    story.append(Spacer(1, 8))

    def make_banner(title_text):
        p = Paragraph(title_text, s_banner)
        t = Table([[p]], colWidths=[566.47])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#1a2b4c')),
            ('TOPPADDING', (0,0), (-1,-1), 3),
            ('BOTTOMPADDING', (0,0), (-1,-1), 3),
            ('LEFTPADDING', (0,0), (-1,-1), 6),
            ('RIGHTPADDING', (0,0), (-1,-1), 6),
        ]))
        return t

    def format_item(item):
        if isinstance(item, dict):
            if 'title' in item and 'content' in item:
                return f"<b>{item['title']}:</b> {item['content']}"
            elif 'myth' in item and 'reality' in item:
                return f"❌ \"{item['myth']}\" → {item['reality']}"
            else:
                return "<br/>".join([f"<b>{k.title()}:</b> {v}" for k, v in item.items()])
        return str(item)

    # 3. Subtopics
    sub_topics = dataset.get('sub_topics', [])
    for idx, st in enumerate(sub_topics, 1):
        st_num = st.get('sub_topic_number', idx)
        st_name = st.get('sub_topic_name', f'SUB TOPIC {st_num}').upper()

        story.append(Paragraph(f"SUB TOPIC {st_num}: {st_name}", s_subtopic_title))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#1a2b4c"), spaceBefore=2, spaceAfter=6))

        # A. Concept
        concept = st.get('concept', '')
        if concept:
            story.append(make_banner("A. Concept"))
            story.append(Paragraph(str(concept), s_body))
            story.append(Spacer(1, 3))

        # B. Prerequisites
        prereqs = st.get('prerequisites', [])
        if prereqs:
            story.append(make_banner("B. Prerequisites"))
            story.append(Paragraph("Before learning this, a student should know:", s_body))
            for p in prereqs:
                p_text = format_item(p)
                if not p_text.startswith("-") and not p_text.startswith("•"):
                    p_text = f"- {p_text}"
                story.append(Paragraph(p_text, s_bullet))
            story.append(Spacer(1, 3))

        # C. Explanation
        exp = st.get('explanation', '')
        if exp:
            story.append(make_banner("C. Explanation"))
            if isinstance(exp, list):
                for paragraph in exp:
                    story.append(Paragraph(format_item(paragraph), s_body))
            else:
                for paragraph in str(exp).split('\n\n'):
                    if paragraph.strip():
                        story.append(Paragraph(paragraph.strip(), s_body))
            story.append(Spacer(1, 3))

        # D. Examples
        examples = st.get('examples', [])
        if examples:
            story.append(make_banner("D. Examples"))
            for ex_i, ex in enumerate(examples, 1):
                ex_text = format_item(ex)
                story.append(Paragraph(ex_text, s_bullet))
            story.append(Spacer(1, 3))

        # E. Practice Problems
        problems = st.get('practice_problems', [])
        if problems:
            story.append(make_banner("E. Practice Problems"))
            for pr_i, pr in enumerate(problems, 1):
                pr_text = format_item(pr)
                story.append(Paragraph(pr_text, s_bullet))
            story.append(Spacer(1, 3))

        # F. Common Misconceptions
        miscs = st.get('common_misconceptions', [])
        if miscs:
            story.append(make_banner("F. Common Misconceptions"))
            for m in miscs:
                m_text = format_item(m)
                story.append(Paragraph(m_text, s_bullet))
            story.append(Spacer(1, 3))

        # G. Assessment
        asmt = st.get('assessment', [])
        if asmt:
            story.append(make_banner("G. Assessment"))
            if isinstance(asmt, list):
                for a_i, a in enumerate(asmt, 1):
                    a_text = format_item(a)
                    # Replace newline characters with html breaks for proper formatting
                    a_html = a_text.replace('\n', '<br/>')
                    story.append(Paragraph(a_html, s_bullet))
                    story.append(Spacer(1, 2))
            elif isinstance(asmt, dict):
                questions = asmt.get('questions', [])
                for q_i, q in enumerate(questions, 1):
                    story.append(Paragraph(f"<b>{q_i}.</b> {format_item(q)}", s_bullet))
            else:
                story.append(Paragraph(str(asmt), s_body))
            story.append(Spacer(1, 5))

    # 4. Short Note & Long Note at end of chapter
    short_note = dataset.get('short_note', '')
    if short_note:
        story.append(make_banner("SHORT NOTE"))
        story.append(Paragraph(str(short_note), s_body))
        story.append(Spacer(1, 5))

    long_note = dataset.get('long_note', '')
    if long_note:
        story.append(make_banner("LONG NOTE"))
        if isinstance(long_note, list):
            for paragraph in long_note:
                story.append(Paragraph(format_item(paragraph), s_body))
        else:
            for paragraph in str(long_note).split('\n\n'):
                if paragraph.strip():
                    story.append(Paragraph(paragraph.strip(), s_body))
        story.append(Spacer(1, 5))

    doc.build(story)
    return buf.getvalue()


def dataset_to_docx(dataset: dict) -> bytes:
    """Convert golden dataset dict to DOCX bytes matching reference structure."""
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('AI Tutor Study Module', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = dataset.get('metadata', {})
    doc.add_paragraph(f"Book Name: {meta.get('book_name', '')}")
    doc.add_paragraph(f"Chapter Name: {meta.get('chapter_name', '')}")
    doc.add_paragraph(f"Chapter Topic: {meta.get('chapter_topic', '')}")

    sub_topics_covered = meta.get('sub_topics_covered', [])
    if isinstance(sub_topics_covered, list):
        sub_topics_str = " | ".join([str(st) for st in sub_topics_covered])
    else:
        sub_topics_str = str(sub_topics_covered)
    doc.add_paragraph(f"Sub Topics Covered: {sub_topics_str}")

    doc.add_paragraph("-" * 50)

    sub_topics = dataset.get('sub_topics', [])
    for idx, st in enumerate(sub_topics, 1):
        st_num = st.get('sub_topic_number', idx)
        st_name = st.get('sub_topic_name', f'SUB TOPIC {st_num}').upper()

        doc.add_heading(f"SUB TOPIC {st_num}: {st_name}", level=1)

        for key, banner_title in [
            ('concept', 'A. Concept'),
            ('prerequisites', 'B. Prerequisites'),
            ('explanation', 'C. Explanation'),
            ('examples', 'D. Examples'),
            ('practice_problems', 'E. Practice Problems'),
            ('common_misconceptions', 'F. Common Misconceptions'),
            ('assessment', 'G. Assessment')
        ]:
            val = st.get(key)
            if val:
                doc.add_heading(banner_title, level=2)
                if isinstance(val, list):
                    for item in val:
                        doc.add_paragraph(str(item), style='List Bullet')
                elif isinstance(val, dict):
                    for k, v in val.items():
                        doc.add_paragraph(f"{k.title()}: {v}")
                else:
                    doc.add_paragraph(str(val))

    short_note = dataset.get('short_note')
    if short_note:
        doc.add_heading('SHORT NOTE', level=1)
        doc.add_paragraph(str(short_note))

    long_note = dataset.get('long_note')
    if long_note:
        doc.add_heading('LONG NOTE', level=1)
        doc.add_paragraph(str(long_note))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
